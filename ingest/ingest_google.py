"""
Re-ingestion script using Google's text-embedding-004 API (new google-genai SDK).
Run this once locally to rebuild the vectorstore/ folder.
This replaces chromadb entirely — no onnxruntime needed on Railway.

Usage:
    py ingest/ingest_google.py
"""
import os
import json
import numpy as np
from pathlib import Path
from docx import Document
from google import genai
from dotenv import load_dotenv

load_dotenv()

client = genai.Client(
    api_key=os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
)

KB_DOCS_DIR = Path('./ingest/kb_docs')
VECTORSTORE_DIR = Path('./vectorstore')
CHUNK_SIZE = 400
CHUNK_OVERLAP = 50
EMBED_MODEL = "models/gemini-embedding-001"


def extract_text_from_docx(filepath: Path):
    doc = Document(filepath)
    text_parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())
    return '\n'.join(text_parts)


def chunk_text(text: str):
    words = text.split()
    chunks = []
    for i in range(0, len(words), CHUNK_SIZE - CHUNK_OVERLAP):
        chunk = ' '.join(words[i:i + CHUNK_SIZE])
        if len(chunk.strip()) > 50:
            chunks.append(chunk)
    return chunks


def embed_texts(texts):
    embeddings = []
    for i, text in enumerate(texts):
        result = client.models.embed_content(
            model=EMBED_MODEL,
            contents=text
        )
        embeddings.append(result.embeddings[0].values)
        if (i + 1) % 10 == 0:
            print(f"  Embedded {i + 1}/{len(texts)} chunks...")
    return embeddings


def run_ingestion():
    VECTORSTORE_DIR.mkdir(exist_ok=True)

    all_documents = []
    all_embeddings = []

    for docx_file in sorted(KB_DOCS_DIR.glob('*.docx')):
        print(f'Processing: {docx_file.name}')
        raw_text = extract_text_from_docx(docx_file)
        chunks = chunk_text(raw_text)
        print(f'  Embedding {len(chunks)} chunks...')
        embeddings = embed_texts(chunks)
        all_documents.extend(chunks)
        all_embeddings.extend(embeddings)

    # Save as lightweight numpy + JSON
    np.save(str(VECTORSTORE_DIR / 'embeddings.npy'), np.array(all_embeddings, dtype=np.float32))
    with open(str(VECTORSTORE_DIR / 'documents.json'), 'w', encoding='utf-8') as f:
        json.dump(all_documents, f, ensure_ascii=False)

    print(f"\nIngestion complete! {len(all_documents)} chunks stored in vectorstore/")
    print(f"Embeddings shape: {np.array(all_embeddings).shape}")


if __name__ == "__main__":
    run_ingestion()
