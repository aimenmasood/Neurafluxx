import os
from pathlib import Path
from docx import Document
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv

load_dotenv()

# CONFIGURATION
KB_DOCS_DIR = Path('./ingest/kb_docs')
CHROMA_DIR = os.getenv('CHROMA_PERSIST_DIR', './chroma_db')
COLLECTION = os.getenv('CHROMA_COLLECTION_NAME', 'neuraflux_kb')
CHUNK_SIZE = 400 
CHUNK_OVERLAP = 50
EMBED_MODEL = 'all-MiniLM-L6-v2'

# LOAD EMBEDDING MODEL
embedder = SentenceTransformer(EMBED_MODEL)

def extract_text_from_docx(filepath: Path):
    doc = Document(filepath)
    text_parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())
    return '\n'.join(text_parts)

def chunk_text(text: str):
    words = text.split()
    chunks = []
    for i in range(0, len(words), CHUNK_SIZE - CHUNK_OVERLAP):
        chunk = ' '.join(words[i : i + CHUNK_SIZE])
        if len(chunk.strip()) > 50:
            chunks.append(chunk)
    return chunks

def run_ingestion():
    client = chromadb.PersistentClient(path=CHROMA_DIR, settings=Settings(anonymized_telemetry=False))
    
    try:
        client.delete_collection(COLLECTION)
    except:
        pass
    
    collection = client.create_collection(name=COLLECTION, metadata={'hnsw:space': 'cosine'})
    
    for docx_file in sorted(KB_DOCS_DIR.glob('*.docx')):
        print(f'Processing: {docx_file.name}')
        raw_text = extract_text_from_docx(docx_file)
        chunks = chunk_text(raw_text)
        
        collection.add(
            ids=[f"{docx_file.stem}_{i}" for i in range(len(chunks))],
            documents=chunks,
            embeddings=embedder.encode(chunks).tolist(),
            metadatas=[{'source': docx_file.stem} for _ in chunks]
        )
    print(f"Ingestion complete. Total chunks stored.")

if __name__ == "__main__":
    run_ingestion()